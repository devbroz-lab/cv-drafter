"""Extract text from .docx files with structural tags for the AI pipeline.

Design notes
------------
The output is a lightly *tagged* plain-text rendering — section markers
(``[HEADING]`` / ``[BOLD]`` / ``[NORMAL]``), table boundaries
(``[TABLE n]`` … rows … ``[END TABLE]``), and header/footer blocks — that gives
Agent 1 structural cues without trying to be a full layout engine.

Robustness is the priority: real-world CVs are produced by Word, Google Docs,
LibreOffice, pandoc, and online builders, which emit structurally valid but odd
``.docx`` files. This module must never crash on such a file — a malformed
paragraph/table is skipped, not fatal — and it captures content Word hides:
headers/footers, text boxes, and hyperlink/field text.
"""

from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

# Qualified tag names (resolved once).
_TR = qn("w:tr")
_TC = qn("w:tc")
_SDT = qn("w:sdt")
_SDT_CONTENT = qn("w:sdtContent")
_P = qn("w:p")
_T = qn("w:t")
_PPR = qn("w:pPr")
_NUMPR = qn("w:numPr")
_TCPR = qn("w:tcPr")
_GRIDSPAN = qn("w:gridSpan")
_VAL = qn("w:val")


def _para_text(p_element) -> str:
    """Concatenate every ``w:t`` descendant of a paragraph element.

    Unlike ``python-docx``'s ``Paragraph.text`` (which only sees runs directly
    under the paragraph), this descends into ``w:hyperlink`` / field runs and
    into anchored text boxes (``w:drawing`` → ``w:txbxContent``), so linked
    emails/URLs and text-box content are not dropped.
    """
    return "".join(t.text for t in p_element.iter(_T) if t.text).strip()


def _is_list_para(p_element) -> bool:
    """True if the paragraph is a numbered/bulleted list item (``w:numPr``)."""
    pPr = p_element.find(_PPR)
    return pPr is not None and pPr.find(_NUMPR) is not None


def _cell_text(tc_element) -> str:
    """Text of one table cell — paragraph lines joined by newline."""
    lines = [_para_text(p) for p in tc_element.iter(_P)]
    return "\n".join(line for line in lines if line)


def _collect_tc_nodes(tr_element) -> list:
    """Return the ``w:tc`` cell elements of a row, including cells wrapped in a
    Word content control (``w:sdt``) that ``python-docx`` hides from ``row.cells``
    (e.g. first-name / family-name fields in some donor templates)."""
    nodes = []
    for child in tr_element.iterchildren():
        if child.tag == _TC:
            nodes.append(child)
        elif child.tag == _SDT:
            for sdt_child in child.iterchildren():
                if sdt_child.tag != _SDT_CONTENT:
                    continue
                for content_child in sdt_child.iterchildren():
                    if content_child.tag == _TC:
                        nodes.append(content_child)
    return nodes


def _grid_span(tc_element) -> int:
    """Horizontal merge width of a cell (``w:gridSpan``), default 1."""
    tcPr = tc_element.find(_TCPR)
    if tcPr is None:
        return 1
    gs = tcPr.find(_GRIDSPAN)
    if gs is None:
        return 1
    try:
        return max(1, int(gs.get(_VAL)))
    except (TypeError, ValueError):
        return 1


def _extract_table_rows(tbl_element) -> list[str]:
    """Extract a table to ``" | "``-joined row strings.

    Column alignment is preserved across merged cells: a horizontally merged
    cell (``w:gridSpan``) is padded with empty placeholder columns, and a
    vertically merged continuation cell is already an empty ``w:tc`` so it keeps
    its column slot. Trailing empty cells are trimmed; fully empty rows are
    dropped.
    """
    rows: list[str] = []
    for tr in tbl_element.iterchildren(tag=_TR):
        cells: list[str] = []
        for tc in _collect_tc_nodes(tr):
            cells.append(_cell_text(tc))
            cells.extend([""] * (_grid_span(tc) - 1))  # keep columns aligned
        # Trim trailing empties (keep internal ones for column alignment).
        while cells and not cells[-1]:
            cells.pop()
        if cells:
            rows.append(" | ".join(cells))
    return rows


def _extract_header_footer(document, which: str) -> list[str]:
    """Collect distinct header or footer blocks across sections.

    Headers/footers commonly carry the candidate name and contact details and
    are *not* part of the document body, so the body walk misses them. Linked
    sections repeat the same content — de-duplicate by block text.
    """
    label = which.upper()
    out: list[str] = []
    seen: set[str] = set()
    for section in document.sections:
        try:
            part = getattr(section, which)  # .header / .footer
            lines: list[str] = []
            for p in part.paragraphs:
                t = _para_text(p._p)
                if t:
                    lines.append(t)
            for tbl in part.tables:
                lines.extend(_extract_table_rows(tbl._tbl))
        except Exception:
            continue
        block = "\n".join(lines).strip()
        if block and block not in seen:
            seen.add(block)
            out.append(f"[{label}]")
            out.extend(lines)
            out.append(f"[/{label}]")
    return out


def extract_text_from_bytes(file_bytes: bytes) -> str:
    if not file_bytes:
        raise ValueError("DOCX payload is empty")

    document = Document(BytesIO(file_bytes))
    chunks: list[str] = []

    # Headers first (often the name/contact block).
    chunks.extend(_extract_header_footer(document, "header"))

    # Body — paragraphs and tables in document order. Each element is isolated
    # so one malformed node can't abort the whole extraction.
    table_index = 0
    for child in document.element.body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        try:
            if tag == "p":
                para = Paragraph(child, document)
                text = _para_text(child)
                if not text:
                    continue
                prefix = "- " if _is_list_para(child) else ""
                style = para.style
                style_name = (style.name if style and style.name else "").lower()
                if "heading" in style_name or "title" in style_name:
                    chunks.append(f"[HEADING] {prefix}{text}")
                elif any(run.bold for run in para.runs if run.text and run.text.strip()):
                    chunks.append(f"[BOLD] {prefix}{text}")
                else:
                    chunks.append(f"[NORMAL] {prefix}{text}")
            elif tag == "tbl":
                table_index += 1
                chunks.append(f"[TABLE {table_index}]")
                chunks.extend(_extract_table_rows(child))
                chunks.append("[END TABLE]")
        except Exception:
            # Per-element resilience: skip a malformed paragraph/table.
            continue

    # Footers last.
    chunks.extend(_extract_header_footer(document, "footer"))

    return "\n".join(chunks).strip()
