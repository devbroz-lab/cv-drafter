"""
Tests for the hardened DOCX extractor (pipeline/extractor/docx_extractor.py):
robustness on odd-but-valid docx, header/footer + hyperlink capture, list
markers, and merged-cell column alignment.
"""

import io
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls

from pipeline.extractor import docx_extractor as dx

SAMPLE = Path(__file__).parent / "sample_files"


def _docx_bytes(document) -> bytes:
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


class TestKeithRegression:
    """Generator-produced docx with no default paragraph style (83/151 paras have
    style==None) previously crashed at `para.style.name`."""

    def test_extracts_without_crashing_and_yields_text(self):
        b = (SAMPLE / "Gen-Keith_Katyora_CV_Detailed.docx").read_bytes()
        t = dx.extract_text_from_bytes(b)
        assert len(t) > 5000
        assert "[TABLE 1]" in t
        assert "Katyora" in t

    def test_header_block_captured(self):
        # The candidate name/contact lives in the header (body walk used to miss it).
        b = (SAMPLE / "Gen-Keith_Katyora_CV_Detailed.docx").read_bytes()
        t = dx.extract_text_from_bytes(b)
        assert "[HEADER]" in t


class TestParaText:
    def test_hyperlink_run_text_captured(self):
        # python-docx's Paragraph.text drops hyperlink runs; _para_text must keep them.
        xml = (
            f'<w:p {nsdecls("w")}>'
            f"<w:hyperlink><w:r><w:t>jane@example.com</w:t></w:r></w:hyperlink>"
            f"<w:r><w:t> profile</w:t></w:r></w:p>"
        )
        p = parse_xml(xml)
        assert dx._para_text(p) == "jane@example.com profile"


class TestHeadersFooters:
    def test_header_and_footer_blocks(self):
        d = Document()
        d.sections[0].header.paragraphs[0].text = "Jane Doe - Senior Expert"
        d.sections[0].footer.paragraphs[0].text = "Confidential footer"
        d.add_paragraph("Body content here.")
        t = dx.extract_text_from_bytes(_docx_bytes(d))
        assert "[HEADER]" in t and "Jane Doe" in t
        assert "[FOOTER]" in t and "Confidential footer" in t
        assert "Body content here." in t


class TestLists:
    def test_direct_numpr_prefixes_dash(self):
        d = Document()
        p = d.add_paragraph("First bullet")
        p._p.get_or_add_pPr().append(OxmlElement("w:numPr"))
        t = dx.extract_text_from_bytes(_docx_bytes(d))
        assert "- First bullet" in t


class TestTables:
    def test_basic_table_rows(self):
        d = Document()
        tbl = d.add_table(rows=2, cols=2)
        tbl.cell(0, 0).text = "Name"
        tbl.cell(0, 1).text = "Jane"
        tbl.cell(1, 0).text = "Role"
        tbl.cell(1, 1).text = "Lead"
        t = dx.extract_text_from_bytes(_docx_bytes(d))
        assert "[TABLE 1]" in t and "[END TABLE]" in t
        assert "Name | Jane" in t
        assert "Role | Lead" in t

    def test_horizontal_merge_keeps_row_alignment(self):
        d = Document()
        tbl = d.add_table(rows=2, cols=3)
        for c, v in enumerate(["A", "B", "C"]):
            tbl.cell(0, c).text = v
        merged = tbl.cell(1, 0).merge(tbl.cell(1, 1))  # gridSpan=2
        merged.text = "MERGED"
        tbl.cell(1, 2).text = "Z"
        t = dx.extract_text_from_bytes(_docx_bytes(d))
        lines = t.splitlines()
        assert "A | B | C" in t
        # The merged cell and the trailing cell stay on the SAME row (alignment kept).
        assert any("MERGED" in ln and "Z" in ln for ln in lines)


class TestEmpty:
    def test_empty_bytes_raises(self):
        import pytest
        with pytest.raises(ValueError):
            dx.extract_text_from_bytes(b"")
